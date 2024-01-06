from Composite.Component import Component

# from Composite.Category import Category
from Database.Database import Database
from State import ArticleStatusContext

from docx import Document

import os

ARTICLES_PATH = str("data")


class Article(Component):
    def __init__(
        self,
        title: str,
        content: str,
        image_path: str,
        status: "ArticleStatusContext",
        meta_description: str,
        creation_date: str,
        parent: Component = None,
    ):
        self.title = title
        self.content = content
        self.image_path = image_path
        self.status = status
        self.meta_description = meta_description
        self.creation_date = creation_date
        self.parent = parent

        self.db = Database(ARTICLES_PATH)

    def operation(self) -> str:
        return f"Article({self.title})"

    def save(self, category_component: Component) -> None:
        # Save article content to DOCX
        article_docx_path = os.path.join(category_component.path, f"{self.title}.docx")
        doc = Document()
        doc.add_heading(self.title, level=1)
        doc.add_paragraph(f"Category: {category_component.name}")
        doc.add_paragraph(f"Status: {self.status}")
        doc.add_paragraph(f"Meta Description: {self.meta_description}")
        doc.add_paragraph(f"Creation Date: {self.creation_date}")
        doc.add_paragraph("\nContent:")
        doc.add_paragraph(self.content)
        doc.save(article_docx_path)

    def to_dict(self) -> dict:
        return {
            "type": "Article",
            "title": self.title,
            "content": self.content,
            "image_path": self.image_path,
            "status": self.status,
            "meta_description": self.meta_description,
            "creation_date": self.creation_date,
            "parent": self.parent,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "Article":
        return cls(
            title=data["title"],
            content=data["content"],
            image_path=data["image_path"],
            status=data["status"],
            meta_description=data["meta_description"],
            creation_date=data["creation_date"],
            parent=data["parent"],
        )
