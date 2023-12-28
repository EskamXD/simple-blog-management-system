from abc import ABC, abstractmethod
from docx import Document
from PIL import ImageFont, ImageDraw, Image

import json
import os
import shutil

ARTICLES_PATH = str("data")

class ArticleState:
    @abstractmethod
    def get_status(self):
        pass

class DraftState(ArticleState):
    @staticmethod
    def get_status():
        return "draft"

class ReadyToPublishState(ArticleState):
    @staticmethod
    def get_status():
        return "ready to publish"

class PublishedState(ArticleState):
    @staticmethod
    def get_status():
        return "published"

class ArticleStatusContext:
    def __init__(self):
        self.state = DraftState().get_status()

    def set_state(self, state):
        self.state = state.get_status()

    def get_status(self):
        return self.state



class Component(ABC):
    @abstractmethod
    def operation(self) -> None:
        pass

    @abstractmethod
    def to_dict(self) -> dict:
        pass

    @abstractmethod
    def from_dict(cls, data: dict) -> "Component":
        pass

class Category(Component):
    def __init__(self, name: str, parent: str | Component) -> None:
        self.name = name
        self.parent = parent
        self.children = []
        self.path = os.path.join(parent, self.name) if name != "data" else "data"
        self.db = Database(ARTICLES_PATH)

    def add(self, component: Component) -> None:
        self.children.append(component)

    def remove(self, component: Component) -> None:
        self.children.remove(component)

    def operation(self) -> str:
        results = [child.operation() for child in self.children]
        return f"Category({self.name}, {', '.join(results)})"

    def to_dict(self) -> dict:
        children_data = [child.to_dict() for child in self.children]
        return {
            "type": "Category",
            "name": self.name,
            "parent": self.parent,
            "children": children_data,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "Category" or "Article":
        try:
            type = data["type"]
        except KeyError:
            raise (KeyError("Missing 'type' key in data dict."))

        if type == "Category":
            name = data["name"]
            parent = data["parent"]
            children = [cls.from_dict(child_data) for child_data in data["children"]]
            category = cls(name, parent)
            category.children = children
            return category
        elif type == "Article":
            return Article.from_dict(data)
        else:
            raise ValueError(f"Invalid type: {type}")

    def save_composite_recursive(self, file_path: str) -> None:
        components_data = self.to_dict()
        self.db.save(file_path, components_data)

    # @classmethod
    # def load_composite_recursive(cls, file_path: str) -> "Category":
    #     try:
    #         data = cls.db.read(file_path)
    #     except FileNotFoundError:
    #         return cls("data")

    #     return cls.from_dict(data)

    def load_composite_recursive(self, file_path: str) -> None:
        try:
            data = self.db.read(file_path)
        except FileNotFoundError:
            return

        self.name = data["name"]
        self.children = [self.from_dict(child_data) for child_data in data["children"]]
        print(self.children)

    def check_existing_title(self, title: str) -> bool:
        for child in self.children:
            if isinstance(child, Article):
                if child.title == title:
                    return True
        return False

class Article(Component):
    def __init__(
        self,
        title: str,
        content: str,
        image_path: str,
        status: "ArticleStatusContext",
        meta_description: str,
        parent: "Category" = None,
    ):
        self.title = title
        self.content = content
        self.image_path = image_path
        self.status = status
        self.meta_description = meta_description
        self.parent = parent

        self.db = Database(ARTICLES_PATH)

    def operation(self) -> str:
        return f"Article({self.title})"

    def save(self, category_component: "Category") -> None:
        # Save article content to DOCX
        article_docx_path = os.path.join(category_component.path, f"{self.title}.docx")
        doc = Document()
        doc.add_heading(self.title, level=1)
        doc.add_paragraph(f"Category: {category_component.name}")
        doc.add_paragraph(f"Status: {self.status}")
        doc.add_paragraph(f"Meta Description: {self.meta_description}")
        doc.add_paragraph("\nContent:")
        doc.add_paragraph(self.content)
        doc.save(article_docx_path)

    def to_dict(self) -> dict:
        return {
            "type": "Article",
            "title": self.title,
            "content": self.content,
            "image_path": self.image_path,
            "status": self.status,
            "meta_description": self.meta_description,
            "parent": self.parent,
        }

    @classmethod
    def from_dict(cls, data: dict) -> "Article":
        return cls(
            title=data["title"],
            content=data["content"],
            image_path=data["image_path"],
            status=data["status"],
            meta_description=data["meta_description"],
            parent=data["parent"],
        )


class Database:
    _instance = None

    def __new__(cls, data_path: str) -> "Database":
        if cls._instance is None:
            cls._instance = super(Database, cls).__new__(cls)
            cls._instance.data_path = data_path
            if not os.path.exists(cls._instance.data_path):
                os.makedirs(cls._instance.data_path)
        return cls._instance

    def save(self, file_path: str, data: dict or list or str) -> None:
        with open(file_path, "w", encoding="utf-8") as json_file:
            json.dump(data, json_file, indent=4)

    def append(self, file_path: str, data: dict or list or str) -> None:
        with open(file_path, "a", encoding="utf-8") as json_file:
            json.dump(data, json_file, indent=4)

    def read(self, file_path: str) -> dict or list or str:
        try:
            with open(file_path, "r", encoding="utf-8") as json_file:
                data = json.load(json_file)
        except FileNotFoundError:
            return FileNotFoundError
        return data
    
    def copy_image(image_path: str, category_component: "Category", title: str) -> None:
            current_path = os.path.dirname(os.path.abspath(__file__))
            new_image_path = os.path.join(current_path, f"{ARTICLES_PATH}/{category_component.name}", f"images/{title}.jpg")

            shutil.copyfile(image_path, new_image_path)

            return new_image_path


class ValidationStrategy(ABC):
    @abstractmethod
    def validate(self, data):
        pass

class MetaTitleValidator(ValidationStrategy):
    def validate(self, meta_title):
        # if meta_title beetwen 50 and 60 characters and length in pixels for arial 22 is up to 580px
        title_length = len(meta_title)
        
        font = ImageFont.truetype("Arial.ttf", 20)
        image = ImageDraw.Draw(Image.new('RGB', (1, 1)))
        title_length_in_pixels = image.textsize(meta_title, font)

        if title_length_in_pixels <= 580:
            if title_length <= 60:
                return True
            
        return Falseś

class MetaDescriptionValidator(ValidationStrategy):
    def validate(self, meta_description):
        # if meta_description up to 158 characters and length in pixels for arial 22 is up to 920px
        title_length = len(meta_description)
        
        font = ImageFont.truetype("Arial.ttf", 14)
        image = ImageDraw.Draw(Image.new('RGB', (1, 1)))
        title_length_in_pixels = image.textsize(meta_description, font)

        if title_length_in_pixels <= 920:
            if title_length <= 158:
                return True
            
        return False

class ImageSizeValidator(ValidationStrategy):
    def validate(self, image_path):
        # if image size is up to 1200x630px
        image = Image.open(image_path)
        width, height = image.size
        weight = os.path.getsize(image_path)
        format = image.format

        # count original image size ratio
        ratio = width/height
        #resize image by ratio to width 1200px
        image = image.resize((1200, int(1200/ratio)))
        #crop image to 630px height from center
        image = image.crop((0, (height-630)/2, width, (height+630)/2))

        if format == "JPEG" or format == "PNG" or format == "JPG":
            #copress to webp
            image.save(image_path, "webp", quality=80)

        if width <= 1200 and height <= 630:
            if weight <= 2*1024*1024:
                return True

        return False      

class ArticleStatusChanger(ValidationStrategy):
    def validate(self, new_status):
        article_status_context = ArticleStatusContext()
        article_status_dict = {
            "D": DraftState(),
            "DRAFT": DraftState(),
            "R": ReadyToPublishState(),
            "READY TO PUBLISH": ReadyToPublishState(),
            "P": PublishedState(),
            "PUBLISHED": PublishedState(),
        }
        new_status = new_status.strip().upper()

        if new_status in article_status_dict.keys():
            article_status_state = article_status_dict[new_status]
            article_status_context.set_state(article_status_state)
        else:
            article_status_context.set_state(DraftState())

        return article_status_context.get_status()


class BlogManagementTerminal:
    def __init__(self, db: "Database", root: "Category") -> None:
        self.db = db
        self.root = root

        # Load state from JSON
        try:
            self.root.load_composite_recursive("main_state.json")
        except FileNotFoundError:
            pass

    def save_and_exit(self) -> None:
        # Save state to JSON
        for component in self.root.children:
            component.save_composite_recursive(f"{ARTICLES_PATH}/{component.name}/articles_info.json")
        self.root.save_composite_recursive("main_state.json")

        exit()

    def add_article(self) -> None:
        # Wyświetl dostępne kategorie
        category_components_dict = {
            category_object: category_object.name
            for category_object in self.root.children
        }
        print(
            f"Available categories: {[category_name for category_name in category_components_dict.values()]}"
        )

        # Pobierz kategorię
        category = self.user_input("Enter category name: ")
        for category_object, category_name in category_components_dict.items():
            if category_name == category:
                category_component = category_object

        # Sprawdź, czy wpisana kategoria istnieje
        if category_component is None:
            print(f"Category '{category}' does not exist.")
            return

        # Wyświetl dostępne tytuły w danej kategorii
        article_titles = [article_object.title for article_object in category_component.children]
        if article_titles:
            print(f"Existing titles in category '{category_component.name}': {article_titles}")
        else:
            print(f"No articles in category '{category_component.name}'.")

        # Pobierz tytuł artykułu
        title = self.user_input("Enter article title: ", MetaTitleValidator())

        # Sprawdź, czy taki tytuł już istnieje w danej kategorii
        if category_component.check_existing_title(title):
            print(f"Article '{title}' already exists in category '{category_component.name}'.")
            return

        # Pobierz treść artykułu
        content = self.user_input("Enter article content (can_be_null=True): ", can_be_null=True)

        # Pobierz ścieżkę do obrazka
        image_path = self.user_input("Enter image path (can_be_null=True): ", None, True)

        # Jeśli ścieżka do obrazka została podana, skopiuj obrazek do folderu z kategorią
        if image_path != "":
            new_image_path = self.db.copy_image(image_path, category_component, title)
        
        ImageSizeValidator().validate(new_image_path)

        # Pobierz status artykułu
        status = self.user_input_status(ArticleStatusChanger())

        # Pobierz meta description
        meta_description = self.user_input("Enter meta description: ", MetaDescriptionValidator())

        # Stwórz obiekt artykułu
        article_component = Article(
            title, content, new_image_path, status, meta_description, category_component.name
        )
        # Dodaj artykuł do kategorii
        category_component.add(article_component)

        # Zapisz stan do DOCX i JSON
        article_component.save(category_component)
        category_component.save_composite_recursive(
            f"{ARTICLES_PATH}/{category_component.name}/articles_info.json"
        )
        self.root.save_composite_recursive("main_state.json")

        print(f"Article '{title}' added to category '{category_component.name}'.")

    def show_articles(self) -> None:
        # Wyświetl dostępne kategorie
        category_components_dict = {
            category_object: category_object.name
            for category_object in self.root.children
        }
        print(
            f"Available categories: {[category_name for category_name in category_components_dict.values()]}"
        )

        # Pobierz kategorię
        category = self.user_input("Enter category name: ")
        for category_object, category_name in category_components_dict.items():
            if category_name == category:
                category_component = category_object

        # Sprawdź, czy wpisana kategoria istnieje
        if category_component is None:
            print(f"Category '{category}' does not exist.")
            return

        # Wyświetl dostępne tytuły w danej kategorii
        article_objects = [article_object for article_object in category_component.children]
        if article_objects:
            print(f"Existing titles in category '{category_component.name}': {article_objects}")
        else:
            print(f"No articles in category '{category_component.name}'.")

    def update_status(self) -> None:
        # Wyświetl dostępne kategorie
        category_components_dict = {
            category_object: category_object.name
            for category_object in self.root.children
        }
        print(
            f"Available categories: {[category_name for category_name in category_components_dict.values()]}"
        )

        # Pobierz kategorię
        category = self.user_input("Enter category name: ")
        for category_object, category_name in category_components_dict.items():
            if category_name == category:
                category_component = category_object

        # Sprawdź, czy wpisana kategoria istnieje
        if category_component is None:
            print(f"Category '{category}' does not exist.")
            return

        # Wyświetl dostępne tytuły w danej kategorii
        article_titles = [article_object.title for article_object in category_component.children]
        if article_titles:
            print(f"Existing titles in category '{category_component.name}': {article_titles}")
        else:
            print(f"No articles in category '{category_component.name}'.")
            return
        
        title = self.user_input("Enter article title: ")

        new_status = self.user_input_status(ArticleStatusChanger())

        articles = [articles for articles in category_component.children]
        for article in articles:
            if article.title == title:
                article.status = new_status
                article.save(category_component)
                break
        
        # Zapisz stan do DOCX i JSON
        category_component.save_composite_recursive(
            f"{ARTICLES_PATH}/{category_component.name}/articles_info.json"
        )
        self.root.save_composite_recursive("main_state.json")

        print(f"Article '{title}' added to category '{category_component.name}'.")

    def show_categories(self) -> None:
        # Wyświetl dostępne kategorie
        category_components_dict = {
            category_object: category_object.name
            for category_object in self.root.children
        }
        print(
            f"Available categories: {list(category_components_dict.values())}"
        )
        return category_components_dict

    def show_titles(self) -> None:
        category_components_dict = self.show_categories()
        titles_list = []

        for category_component, category_name in category_components_dict.items():
            titles_list.append([article_object.title for article_object in category_component.children])

        print(f"All titles list: {titles_list}")
        return titles_list

    def user_input(self, input_placeholder: str, validator_function: "ValidationStrategy" = None, can_be_null: bool = False) -> str:
        user_input = input(input_placeholder)
        while type(user_input) == str and len(user_input) < 1 and not can_be_null:
            validator_function.validate(user_input)
            user_input = input(input_placeholder)

        return user_input.strip()

    def user_input_status(self, status_validator: "ValidationStrategy") -> str:
        user_input = ""
        while user_input == "":
            # while type(user_input) == str and len(user_input) < 1:
            user_input = input(f"Enter article status {self.article_status_list}: ")
            user_input = status_validator.validate(user_input)

        return user_input

    def print_composite(self) -> None:
        print(self.root.operation())


class BlogManagementGUI:
    def __init__(self, db: "Database", root: "Category") -> None:
        self.db = db
        self.root = root

        # Load state from JSON
        try:
            self.root.load_composite_recursive("main_state.json")
        except FileNotFoundError:
            pass

    def run(self) -> None:
        pass

    def save_and_exit(self) -> None:
        pass

    def add_article(self) -> None:
        pass

    def show_articles(self) -> None:
        pass

    def update_status(self) -> None:
        pass

    def show_categories(self) -> None:
        pass

    def show_titles(self) -> None:
        pass

    def print_composite(self) -> None:
        pass


if __name__ == "__main__":
    blog_management = BlogManagementTerminal(
        Database(ARTICLES_PATH), Category("data", "data")
    )

    choice_dict = {
        "1": blog_management.add_article,
        "2": blog_management.show_articles,
        "3": blog_management.update_status,
        "4": blog_management.show_categories,
        "5": blog_management.show_titles,
        "6": blog_management.print_composite,
        "-1": blog_management.save_and_exit,
    }

    while True:
        print("\nMenu:")
        print("1. Add article")
        print("2. Show articles")
        print("3. Update status")
        print("4. Show categories")
        print("5. Show titles")
        print("6. Print composite")
        print("-1. Save & Exit")

        choice = input("Enter your choice (1-6): ")

        choice_dict.get(choice, lambda: print("Invalid choice. Please try again."))()
